from io import BytesIO
from sqlalchemy.orm import Session
from sqlalchemy import desc

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from app.models.output_document import OutputDocument, OutputSection, OutputSectionVersion


def export_csr_to_docx(document_id: int, db: Session) -> bytes:
    """
    Export OutputDocument to DOCX format.
    
    Loads OutputDocument and related OutputSection + latest OutputSectionVersion for each section.
    Creates a DOCX document with:
    - Document title as main heading
    - Section titles as level 1 headings
    - Section text as paragraphs
    
    Args:
        document_id: ID of the CSR document to export
        db: Database session
        
    Returns:
        bytes: The DOCX file content as bytes
        
    Raises:
        ValueError: If document not found
    """
    # Load OutputDocument
    document = db.query(OutputDocument).filter(OutputDocument.id == document_id).first()
    if not document:
        raise ValueError(f"CSR document with id {document_id} not found")
    
    # Create a new Document
    doc = Document()
    
    # Add document title as main heading
    title_paragraph = doc.add_heading(document.title, level=0)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Load sections ordered by order_index
    sections = (
        db.query(OutputSection)
        .filter(OutputSection.document_id == document_id)
        .order_by(OutputSection.order_index)
        .all()
    )
    
    # Process each section
    for section in sections:
        # Add section title as level 1 heading
        doc.add_heading(section.title, level=1)
        
        # Get the latest version for this section
        latest_version = (
            db.query(OutputSectionVersion)
            .filter(OutputSectionVersion.section_id == section.id)
            .order_by(desc(OutputSectionVersion.created_at), desc(OutputSectionVersion.id))
            .first()
        )
        
        # Add section text if version exists
        if latest_version and latest_version.text:
            # Split text by newlines and add as paragraphs
            text_lines = latest_version.text.strip().split('\n')
            for line in text_lines:
                if line.strip():  # Only add non-empty lines
                    paragraph = doc.add_paragraph(line.strip())
                    # Set paragraph style to normal (default)
                    paragraph.style = doc.styles['Normal']
        else:
            # Add placeholder if no version exists
            doc.add_paragraph("[No content available for this section]")
    
    # Save document to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

