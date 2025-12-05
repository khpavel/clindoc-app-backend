from pathlib import Path
from typing import Optional


def extract_text_from_file(path: Path) -> str:
    """
    Extract text from various file formats.
    Supports:
    - Text files: .txt, .md
    - PDF files: .pdf
    - Word documents: .docx, .doc
    - Excel files: .xlsx, .xls
    """
    suffix = path.suffix.lower()
    
    # Plain text files
    if suffix in [".txt", ".md"]:
        with open(path, encoding="utf-8", errors="ignore") as f:
            return f.read()
    
    # PDF files
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            text_parts = []
            for page in reader.pages:
                text_parts.append(page.extract_text())
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    # Word documents (.docx)
    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(path)
            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    # Word documents (.doc) - old format
    if suffix == ".doc":
        try:
            # Try using python-docx first (may not work for old .doc)
            from docx import Document
            doc = Document(path)
            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            return "\n".join(text_parts)
        except Exception:
            # If python-docx fails, return error message
            raise ValueError(
                "Failed to extract text from .doc file. "
                "Old .doc format requires additional libraries. "
                "Please convert to .docx format or use a tool like antiword."
            )
    
    # Excel files (.xlsx)
    if suffix == ".xlsx":
        try:
            from openpyxl import load_workbook
            workbook = load_workbook(path, data_only=True)
            text_parts = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip():
                        text_parts.append(row_text)
                text_parts.append("")  # Empty line between sheets
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from Excel file: {str(e)}")
    
    # Excel files (.xls) - old format
    if suffix == ".xls":
        try:
            import xlrd
            workbook = xlrd.open_workbook(path)
            text_parts = []
            for sheet_name in workbook.sheet_names():
                sheet = workbook.sheet_by_name(sheet_name)
                text_parts.append(f"Sheet: {sheet_name}")
                for row_idx in range(sheet.nrows):
                    row_values = [str(sheet.cell_value(row_idx, col_idx)) 
                                 for col_idx in range(sheet.ncols)]
                    row_text = "\t".join(row_values)
                    if row_text.strip():
                        text_parts.append(row_text)
                text_parts.append("")  # Empty line between sheets
            return "\n".join(text_parts)
        except ImportError:
            raise ValueError(
                "Failed to extract text from .xls file. "
                "xlrd library is required for old .xls format. "
                "Please install it: pip install xlrd"
            )
        except Exception as e:
            raise ValueError(f"Failed to extract text from Excel file: {str(e)}")
    
    # For unsupported file types
    return f"[[UNPARSED FILE: {path.name} – file type '{suffix}' is not supported]]"

